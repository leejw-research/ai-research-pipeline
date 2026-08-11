import streamlit as st
import os
import json
import time
import io
import sqlite3
import requests
from bs4 import BeautifulSoup
from google import genai
from google.genai import types
from anthropic import Anthropic
import docx

# 페이지 기본 설정
st.set_page_config(
    page_title="Dual-AI Cross Research Pipeline",
    page_icon="🔬",
    layout="wide"
)

CONFIG_FILE = "user_config.json"

# ==========================================
# 0. 데이터베이스(History) 초기화 함수
# ==========================================
def init_db():
    conn = sqlite3.connect("history.db", check_same_thread=False)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS history 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, topic TEXT, result TEXT, timestamp DATETIME DEFAULT CURRENT_TIMESTAMP)''')
    conn.commit()
    conn.close()

init_db()

# ==========================================
# 1. 서비스 전용 접속 비밀번호 설정
# ==========================================
ACCESS_PASSWORD = "wqlab"  # <-- 필요시 원하는 비밀번호로 변경하세요

if "is_master_auth" not in st.session_state:
    st.session_state.is_master_auth = False

if not st.session_state.is_master_auth:
    st.title("🔒 제한된 연구 분석 파이프라인")
    st.markdown("이 프로그램은 승인된 사용자만 이용할 수 있습니다. 접속 비밀번호를 입력해 주세요.")
    
    pw_input = st.text_input("서비스 접속 비밀번호", type="password", placeholder="비밀번호를 입력하세요")
    if st.button("🔐 접속하기", type="primary"):
        if pw_input == ACCESS_PASSWORD:
            st.session_state.is_master_auth = True
            st.rerun()
        else:
            st.error("❌ 비밀번호가 올바르지 않습니다.")
    st.stop()

# ==========================================
# 2. 설정 및 로그인 관련 함수
# ==========================================
def load_config():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}

def save_config(config_data):
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(config_data, f, ensure_ascii=False, indent=4)

def clear_config():
    if os.path.exists(CONFIG_FILE):
        os.remove(CONFIG_FILE)

saved_cfg = load_config()

if "is_logged_in" not in st.session_state:
    st.session_state.is_logged_in = False

if "gemini_api_key" not in st.session_state:
    st.session_state.gemini_api_key = ""
if "claude_api_key" not in st.session_state:
    st.session_state.claude_api_key = ""

if "remember_info" not in st.session_state:
    st.session_state.remember_info = saved_cfg.get("remember_info", False)
if "gemini_model" not in st.session_state:
    st.session_state.gemini_model = saved_cfg.get("gemini_model", "")
if "claude_model" not in st.session_state:
    st.session_state.claude_model = saved_cfg.get("claude_model", "")

if "final_output" not in st.session_state:
    st.session_state.final_output = ""

# ==========================================
# 2-1. 출력 토큰 한도 설정 (잘림 방지 핵심)
# ==========================================
# 필요에 따라 늘릴 수 있습니다. Claude Sonnet 5는 최대 128000까지 지원합니다.
MAX_OUTPUT_TOKENS_GEMINI = 32000
MAX_OUTPUT_TOKENS_CLAUDE = 32000
MAX_CONTINUATION_ROUNDS = 3  # 한 번 더 잘려도 최대 몇 번까지 이어받을지

# URL 텍스트 추출 함수
def get_url_text(url):
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(response.content, 'html.parser')
        for script in soup(["script", "style"]):
            script.extract()
        return soup.get_text(separator='\n', strip=True)
    except Exception as e:
        return f"[URL 읽기 실패: {e}]"

def create_word_document(text_content):
    doc = docx.Document()
    doc.add_heading("AI Dual-Agent 교차 연구 검증 결과", level=1)
    
    for line in text_content.split("\n"):
        if line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)
        else:
            if line.strip():
                doc.add_paragraph(line)
            else:
                doc.add_paragraph()
                
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ==========================================
# 3. 사이드바 (기록 및 설정 제어)
# ==========================================
with st.sidebar:
    if st.session_state.is_logged_in:
        st.title("📂 내 연구 기록 (History)")
        try:
            conn = sqlite3.connect("history.db", check_same_thread=False)
            c = conn.cursor()
            c.execute("SELECT id, topic FROM history ORDER BY timestamp DESC")
            records = c.fetchall()
            conn.close()
            
            if records:
                for rec in records:
                    display_title = rec[1] if len(rec[1]) <= 18 else rec[1][:18] + "..."
                    if st.button(f"📖 {display_title}", key=f"hist_{rec[0]}", use_container_width=True):
                        conn = sqlite3.connect("history.db", check_same_thread=False)
                        c = conn.cursor()
                        c.execute("SELECT result FROM history WHERE id=?", (rec[0],))
                        res_row = c.fetchone()
                        conn.close()
                        if res_row:
                            st.session_state.final_output = res_row[0]
                            st.rerun()
            else:
                st.info("저장된 연구 기록이 없습니다.")
        except Exception as e:
            st.warning("기록을 불러오는 중 오류가 발생했습니다.")

        st.markdown("---")
        st.success("✅ 로그인 완료")
        st.info(f"📌 **Gemini:** `{st.session_state.gemini_model}`\n\n📌 **Claude:** `{st.session_state.claude_model}`")
        
        if st.button("🔓 로그아웃 / 설정 변경", use_container_width=True):
            st.session_state.is_logged_in = False
            st.session_state.gemini_api_key = ""
            st.session_state.claude_api_key = ""
            st.session_state.final_output = ""
            st.rerun()
    else:
        st.subheader("1. API 키 입력")
        gemini_key_input = st.text_input("Gemini API 키", type="password", placeholder="API 키 입력", key="gemini_key_box")
        claude_key_input = st.text_input("Claude API 키", type="password", placeholder="API 키 입력", key="claude_key_box")
        
        st.markdown("🔗 **API 키 발급 바로가기**")
        col_l1, col_l2 = st.columns(2)
        with col_l1:
            st.link_button("Google", "https://aistudio.google.com/app/apikey", use_container_width=True)
        with col_l2:
            st.link_button("Anthropic", "https://console.anthropic.com/", use_container_width=True)

        st.markdown("---")
        st.subheader("2. 모델 설정")
        gemini_mod_input = st.text_input("Gemini 모델명", value=st.session_state.gemini_model, placeholder="예: gemini-2.5-flash", key="gemini_mod_box")
        claude_mod_input = st.text_input("Claude 모델명", value=st.session_state.claude_model, placeholder="예: claude-sonnet-5", key="claude_mod_box")

        remember_info_toggle = st.toggle("💾 설정 정보 저장", value=st.session_state.remember_info)

        if st.button("🔑 로그인 / 설정 완료", type="primary", use_container_width=True):
            if not gemini_key_input or not claude_key_input:
                st.error("API 키를 모두 입력해 주세요.")
            elif not gemini_mod_input or not claude_mod_input:
                st.error("모델명을 모두 입력해 주세요.")
            else:
                st.session_state.gemini_api_key = gemini_key_input
                st.session_state.claude_api_key = claude_key_input
                st.session_state.gemini_model = gemini_mod_input
                st.session_state.claude_model = claude_mod_input
                st.session_state.remember_info = remember_info_toggle
                st.session_state.is_logged_in = True
                
                if remember_info_toggle:
                    save_config({
                        "remember_info": True,
                        "gemini_model": gemini_mod_input,
                        "claude_model": claude_mod_input
                    })
                else:
                    clear_config()
                st.rerun()

# ==========================================
# 4. 메인 인터페이스
# ==========================================
st.title("🔬 AI Dual-Agent 교차 연구 검증 파이프라인")

tab_main, tab_admin, tab_troubleshoot = st.tabs(["🚀 연구 실행", "🛠️ 관리자 탭", "🚨 예외 대처 가이드"])

with tab_main:
    if not st.session_state.is_logged_in:
        st.warning("👈 왼쪽 사이드바에서 로그인 후 이용해 주세요.")
    else:
        st.subheader("1. 연구 주제 및 참고 자료 입력")
        
        col_t1, col_t2 = st.columns([2, 1])
        with col_t1:
            user_topic = st.text_area("연구 주제 / 핵심 분석 가설 입력", height=100)
        with col_t2:
            url_input = st.text_input("참조할 웹페이지 URL (선택)", placeholder="https://...")
            uploaded_file = st.file_uploader("참고 문서 (.docx)", type=["docx"])
        
        start_btn = st.button("🔥 파이프라인 실행 시작", type="primary", use_container_width=True)

        if start_btn:
            if not user_topic:
                st.warning("연구 주제를 입력해 주세요.")
            else:
                status_container = st.status("🚀 결과 도출 중... 파이프라인을 가동합니다.", expanded=True)
                
                gemini_client = genai.Client(api_key=st.session_state.gemini_api_key)
                claude_client = Anthropic(api_key=st.session_state.claude_api_key)

                # ------------------------------------------------
                # Gemini 호출 (잘림 감지 + 자동 이어받기)
                # ------------------------------------------------
                def safe_call_gemini(contents_payload, sys_prompt):
                    for attempt in range(1, 6):
                        try:
                            full_text = ""
                            current_contents = contents_payload
                            for round_i in range(MAX_CONTINUATION_ROUNDS):
                                res = gemini_client.models.generate_content(
                                    model=st.session_state.gemini_model,
                                    contents=current_contents,
                                    config=types.GenerateContentConfig(
                                        system_instruction=sys_prompt,
                                        temperature=0.7,
                                        max_output_tokens=MAX_OUTPUT_TOKENS_GEMINI,
                                    )
                                )
                                chunk_text = res.text or ""
                                full_text += chunk_text

                                finish_reason = None
                                if getattr(res, "candidates", None):
                                    finish_reason = str(res.candidates[0].finish_reason)

                                if finish_reason == "MAX_TOKENS":
                                    status_container.write(f"✂️ Gemini 응답이 길어 이어서 생성 중... ({round_i + 1}회차)")
                                    current_contents = (
                                        f"{contents_payload}\n\n"
                                        f"[이전까지 작성된 내용]\n{full_text}\n\n"
                                        f"위 내용에 이어서, 끊긴 부분부터 자연스럽게 계속 작성해줘. "
                                        f"이미 작성된 부분은 반복하지 마."
                                    )
                                    continue
                                else:
                                    break

                            return full_text
                        except Exception as e:
                            if any(k in str(e) for k in ["503", "429", "UNAVAILABLE", "Overloaded"]):
                                time.sleep(5)
                            else:
                                raise e
                    raise Exception("Gemini API 재시도 횟수 초과")

                # ------------------------------------------------
                # Claude 호출 (잘림 감지 + 자동 이어받기)
                # ------------------------------------------------
                def safe_call_claude(prompt, sys_prompt):
                    for attempt in range(1, 6):
                        try:
                            messages = [{"role": "user", "content": prompt}]
                            full_text = ""

                            for round_i in range(MAX_CONTINUATION_ROUNDS):
                                res = claude_client.messages.create(
                                    model=st.session_state.claude_model,
                                    max_tokens=MAX_OUTPUT_TOKENS_CLAUDE,
                                    system=sys_prompt,
                                    messages=messages
                                )
                                chunk_text = "\n".join(
                                    [b.text for b in res.content if getattr(b, 'type', None) == 'text']
                                )
                                full_text += chunk_text

                                if res.stop_reason == "max_tokens":
                                    status_container.write(f"✂️ Claude 응답이 길어 이어서 생성 중... ({round_i + 1}회차)")
                                    messages = [
                                        {"role": "user", "content": prompt},
                                        {"role": "assistant", "content": full_text},
                                        {"role": "user", "content": "이어서 계속 작성해줘. 이미 작성된 부분은 반복하지 마."}
                                    ]
                                    continue
                                else:
                                    break

                            return full_text
                        except Exception as e:
                            if any(k in str(e) for k in ["429", "529", "overloaded"]):
                                time.sleep(5)
                            else:
                                raise e
                    raise Exception("Claude API 재시도 횟수 초과")

                try:
                    file_text = ""
                    if uploaded_file is not None:
                        status_container.write("📄 업로드된 문서 분석 중...")
                        doc = docx.Document(uploaded_file)
                        file_text += "\n".join([p.text for p in doc.paragraphs])

                    if url_input:
                        status_container.write(f"🌐 URL 크롤링 중: {url_input}")
                        url_content = get_url_text(url_input)
                        file_text += f"\n\n[참조 웹페이지 내용]\n{url_content}"

                    status_container.write("💡 [1/4] Gemini: 초기 연구 아이디어 및 가설 생성 중...")
                    gemini_prompt = f"주제: {user_topic}\n\n[참고 자료 내용]\n{file_text[:6000]}"
                    gemini_res = safe_call_gemini(gemini_prompt, "학술 연구 기획 전문가로서 독창적 가설과 방법론을 제안하세요.")
                    
                    status_container.write("🔍 [2/4] Claude: 메커니즘 분석 및 학술적 허점 교차 고찰 중...")
                    claude_prompt = f"Gemini의 다음 제안에 대해 학술적 맹점 및 비판적 고찰을 수행하세요:\n\n{gemini_res}"
                    claude_res = safe_call_claude(claude_prompt, "상호 심의 피어 리뷰어로서 독창적이고 날카로운 비판을 제공하세요.")

                    status_container.write("🤝 [3/4] 토론 통합 및 최종 합성 중...")
                    st.session_state.final_output = f"## 1. Gemini 제안\n{gemini_res}\n\n## 2. Claude 교차 검증 및 비판\n{claude_res}"

                    # SQLite 데이터베이스에 기록 저장
                    conn = sqlite3.connect("history.db", check_same_thread=False)
                    c = conn.cursor()
                    c.execute("INSERT INTO history (topic, result) VALUES (?, ?)", (user_topic, st.session_state.final_output))
                    conn.commit()
                    conn.close()

                    status_container.update(label="✅ 분석 완료 및 기록 저장 완료!", state="complete", expanded=False)
                    st.rerun()

                except Exception as pipeline_err:
                    status_container.update(label="💥 파이프라인 중단됨", state="error")
                    st.error(f"오류 발생: {pipeline_err}")

        if st.session_state.final_output:
            st.markdown("---")
            st.success("🎉 분석 결과")
            
            col_dl, col_info = st.columns([1, 2])
            with col_dl:
                word_file = create_word_document(st.session_state.final_output)
                st.download_button(
                    label="📥 워드 파일로 다운로드 (.docx)",
                    data=word_file,
                    file_name="research_analysis_result.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            with col_info:
                st.info("💡 좌측 사이드바에서 이전 연구 기록을 언제든 다시 불러올 수 있습니다.")

            st.markdown(st.session_state.final_output)
            st.markdown("### 📋 전체 텍스트 복사 영역")
            st.code(st.session_state.final_output, language="markdown")

with tab_admin:
    st.subheader("🛠️ 시스템 현황 및 API 실시간 테스트")
    if not st.session_state.is_logged_in:
        st.warning("로그인 후 이용하실 수 있습니다.")
    else:
        col_a1, col_a2 = st.columns(2)
        with col_a1:
            if st.button("🧪 Gemini API 연결 검증"):
                try:
                    c = genai.Client(api_key=st.session_state.gemini_api_key)
                    models = [m.name for m in c.models.list()]
                    st.success(f"Gemini 연결 성공! (모델 수: {len(models)}개)")
                except Exception as e:
                    st.error(f"실패: {e}")
        with col_a2:
            if st.button("🧪 Claude API 연결 검증"):
                try:
                    c = Anthropic(api_key=st.session_state.claude_api_key)
                    models = [m.id for m in c.models.list()]
                    st.success(f"Claude 연결 성공! (모델 수: {len(models)}개)")
                except Exception as e:
                    st.error(f"실패: {e}")

with tab_troubleshoot:
    st.subheader("🚨 예기치 못한 오류 대처 가이드")
    st.markdown("""
    | 에러 종류 | 주요 원인 | 즉각적인 해결 대처방안 |
    | :--- | :--- | :--- |
    | **`404 NOT_FOUND`** | 잘못된 모델명 입력 | 사이드바 재설정 후 정확한 모델명 입력 |
    | **`429 / RESOURCE_EXHAUSTED`** | 호출 한도 초과 | 자동으로 대기 후 재시도 수행 |
    | **`503 / UNAVAILABLE`** | 서버 측 과부하 | 10~30초 후 다시 실행 버튼 클릭 |
    | **응답이 중간에 잘림** | 출력 토큰 한도(`max_tokens`) 초과 | 이번 수정본은 자동으로 감지해 이어서 생성합니다 (최대 3회) |
    """)
